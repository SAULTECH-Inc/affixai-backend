"""Document Processing Layer — Phase B core engine.

The unified conversion + extraction + OCR + normalization API every other
layer can lean on:

  * Integration layer downloads documents in arbitrary formats via these.
  * Workflow layer extracts plain text for indexing / search.
  * Collaboration layer pulls text for diff / preview.

We deliberately do NOT do any storage / DB writes here — this module is
purely format-in / format-out so it stays easy to test and replace.

Format aliases supported:
  - "pdf"  (application/pdf)
  - "docx" (officedocument)
  - "txt"  (plain text)
  - "md"   (markdown)

Conversion matrix is documented in convert_document(); anything missing
falls back to going through PDF as the lingua franca where it makes sense.
"""
from __future__ import annotations

import io
import re
import subprocess
import tempfile
from enum import Enum
from pathlib import Path
from typing import Literal

from loguru import logger


class DocFormat(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"


SupportedFormat = Literal["pdf", "docx", "txt", "md"]


# Mime types we accept on inputs / set on outputs.
MIME_BY_FORMAT: dict[str, str] = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain",
    "md": "text/markdown",
}


def detect_format(filename: str | None, mime_type: str | None) -> DocFormat | None:
    """Guess the source format from filename + MIME. Returns None if unknown."""
    if mime_type:
        m = mime_type.lower()
        if "pdf" in m:
            return DocFormat.PDF
        if "wordprocessingml" in m or "msword" in m:
            return DocFormat.DOCX
        if "markdown" in m:
            return DocFormat.MD
        if m.startswith("text/"):
            return DocFormat.TXT
    if filename:
        ext = Path(filename).suffix.lower().lstrip(".")
        if ext in ("pdf",):
            return DocFormat.PDF
        if ext in ("docx", "doc"):
            return DocFormat.DOCX
        if ext in ("md", "markdown"):
            return DocFormat.MD
        if ext in ("txt", "text", "log"):
            return DocFormat.TXT
    return None


# ---- Normalization ---------------------------------------------------------


_BOM_RE = re.compile(r"^﻿")
_TRIPLE_BLANK_RE = re.compile(r"\n{3,}")


def normalize_text(text: str) -> str:
    """Standardize whitespace + line endings; strip BOMs.

    Idempotent. Used at the OUTPUT side of every extract_* function so callers
    can rely on `\\n` line endings and no triple-blank-line gaps.
    """
    if not text:
        return ""
    s = _BOM_RE.sub("", text)
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = _TRIPLE_BLANK_RE.sub("\n\n", s)
    # Trim trailing whitespace per line — common artifact of PDF extraction.
    s = "\n".join(line.rstrip() for line in s.split("\n"))
    return s.strip() + "\n" if s.strip() else ""


# ---- Text extraction (with OCR fallback) -----------------------------------


def extract_text_from_pdf(pdf_bytes: bytes, *, ocr_fallback: bool = True) -> str:
    """Extract plain text from a PDF.

    If the PDF is digital-native (has selectable text), PyMuPDF gives us
    the text directly. If a page returns empty text and `ocr_fallback` is
    True, we rasterize that page and run Tesseract on it.

    The OCR fallback handles scanned PDFs transparently — callers don't
    need to know whether the source was digital or scanned.
    """
    import fitz

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as exc:
        logger.warning(f"pdf open failed: {exc}")
        return ""

    pages_out: list[str] = []
    for i, page in enumerate(doc):
        text = page.get_text("text") or ""
        if not text.strip() and ocr_fallback:
            ocr_text = _ocr_page(page, page_index=i)
            if ocr_text:
                text = ocr_text
        pages_out.append(text)

    doc.close()
    return normalize_text("\n\n".join(pages_out))


def _ocr_page(page, page_index: int) -> str:
    """Rasterize a single PyMuPDF page and run Tesseract.

    We render at 2x (200dpi) — sweet spot between accuracy and speed.
    """
    try:
        import pytesseract  # type: ignore[import-untyped]
        from PIL import Image
    except ImportError:
        logger.warning("OCR libraries not installed — skipping fallback")
        return ""

    try:
        pix = page.get_pixmap(matrix=__import__("fitz").Matrix(2, 2))
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        text = pytesseract.image_to_string(img) or ""
        if text.strip():
            logger.info(f"OCR fallback recovered text on page {page_index}")
        return text
    except Exception as exc:
        logger.warning(f"OCR failed for page {page_index}: {exc}")
        return ""


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Plain text from a DOCX. Preserves paragraph + table structure roughly."""
    from docx import Document  # python-docx

    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        logger.warning(f"docx open failed: {exc}")
        return ""

    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text]
            if cells:
                parts.append(" | ".join(cells))
    return normalize_text("\n".join(parts))


# ---- Markdown extraction ---------------------------------------------------


def pdf_to_markdown(pdf_bytes: bytes) -> str:
    """Render a PDF as Markdown.

    PyMuPDF doesn't have a first-party "extract markdown" mode, so we walk
    the page dict and synthesize a markdown structure:
      * heading-sized spans → `#` / `##` based on relative font size
      * everything else → flowing paragraphs with blank lines between blocks
      * bullet markers and numbered lists are preserved as-is

    This is good enough for downstream rendering / preview; not lossless.
    """
    import fitz

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as exc:
        logger.warning(f"pdf open failed: {exc}")
        return ""

    # Pass 1: collect every span's size to figure out the heading thresholds.
    sizes: list[float] = []
    for page in doc:
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    sizes.append(float(span.get("size") or 10))
    if not sizes:
        doc.close()
        return ""
    body_size = _mode(sizes)
    # Heuristic: anything ≥1.4x body is a heading; ≥1.8x is H1.
    h1_threshold = body_size * 1.8
    h2_threshold = body_size * 1.4

    parts: list[str] = []
    for page in doc:
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:
                continue
            block_lines: list[str] = []
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                text = "".join(s.get("text", "") for s in spans).strip()
                if not text:
                    continue
                line_size = max(
                    (float(s.get("size") or 10) for s in spans), default=10
                )
                if line_size >= h1_threshold:
                    block_lines.append(f"# {text}")
                elif line_size >= h2_threshold:
                    block_lines.append(f"## {text}")
                else:
                    block_lines.append(text)
            if block_lines:
                parts.append("\n".join(block_lines))
    doc.close()
    return normalize_text("\n\n".join(parts))


def docx_to_markdown(docx_bytes: bytes) -> str:
    """Render a DOCX as Markdown. Maps Word heading styles → ATX headings."""
    from docx import Document

    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        logger.warning(f"docx open failed: {exc}")
        return ""

    parts: list[str] = []
    for p in doc.paragraphs:
        text = p.text
        if not text:
            parts.append("")
            continue
        style = (p.style.name or "").lower() if p.style else ""
        if style.startswith("heading 1") or style == "title":
            parts.append(f"# {text}")
        elif style.startswith("heading 2"):
            parts.append(f"## {text}")
        elif style.startswith("heading 3"):
            parts.append(f"### {text}")
        elif style == "list paragraph":
            parts.append(f"- {text}")
        else:
            parts.append(text)

    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        parts.append("| " + " | ".join(header_cells) + " |")
        parts.append("| " + " | ".join("---" for _ in header_cells) + " |")
        for row in table.rows[1:]:
            parts.append("| " + " | ".join(c.text.strip() for c in row.cells) + " |")
        parts.append("")
    return normalize_text("\n\n".join(parts))


# ---- DOCX generation -------------------------------------------------------


def pdf_to_docx(pdf_bytes: bytes) -> bytes:
    """Convert PDF → DOCX via the pdf2docx library.

    Quality is decent for digital PDFs; scanned PDFs come out as a single
    block per page (effectively text + image layers) — callers wanting clean
    DOCX from a scanned source should OCR first and use `text_to_docx`.
    """
    from pdf2docx import Converter

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as src:
        src.write(pdf_bytes)
        src_path = src.name
    out_path = src_path.replace(".pdf", ".docx")
    try:
        cv = Converter(src_path)
        cv.convert(out_path, start=0, end=None)
        cv.close()
        return Path(out_path).read_bytes()
    finally:
        Path(src_path).unlink(missing_ok=True)
        Path(out_path).unlink(missing_ok=True)


def text_to_docx(text: str, *, title: str | None = None) -> bytes:
    """Create a minimally-styled DOCX from plain text.

    Each blank-line-separated block becomes a paragraph. Markdown-ish heading
    cues (`# …`, `## …`) become Word heading styles so the doc reads well in
    Word / Pages / Google Docs.
    """
    from docx import Document

    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    for block in (text or "").split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            doc.add_heading(block[2:].strip(), level=1)
        elif block.startswith("## "):
            doc.add_heading(block[3:].strip(), level=2)
        elif block.startswith("### "):
            doc.add_heading(block[4:].strip(), level=3)
        else:
            doc.add_paragraph(block)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def markdown_to_docx(md_text: str) -> bytes:
    """Markdown → DOCX. We re-use text_to_docx since it already understands
    `#`/`##`/`###` heading cues — Markdown's headings are the most
    consequential structure for our use case."""
    return text_to_docx(md_text)


# ---- PDF generation --------------------------------------------------------


def text_to_pdf(text: str, *, title: str | None = None) -> bytes:
    """Plain text → PDF via ReportLab. Single column, Letter-sized, 1in margins.

    For headings, Markdown-ish `# `/`## `/`### ` prefixes are honoured so
    Markdown-source documents round-trip somewhat cleanly.
    """
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    styles = getSampleStyleSheet()
    body = styles["BodyText"]
    body.leading = 14
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], spaceBefore=12, spaceAfter=6)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], spaceBefore=10, spaceAfter=5)
    h3 = ParagraphStyle("h3", parent=styles["Heading3"], spaceBefore=8, spaceAfter=4)

    flowables: list = []
    if title:
        flowables.append(Paragraph(_html_escape(title), styles["Title"]))
        flowables.append(Spacer(1, 12))
    for block in (text or "").split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            flowables.append(Paragraph(_html_escape(block[2:]), h1))
        elif block.startswith("## "):
            flowables.append(Paragraph(_html_escape(block[3:]), h2))
        elif block.startswith("### "):
            flowables.append(Paragraph(_html_escape(block[4:]), h3))
        else:
            # Preserve inline newlines so soft-broken text doesn't collapse
            # into a single visual blob.
            for line in block.split("\n"):
                flowables.append(Paragraph(_html_escape(line), body))
        flowables.append(Spacer(1, 6))

    doc.build(flowables)
    return buf.getvalue()


def markdown_to_pdf(md_text: str) -> bytes:
    """Markdown → PDF. Same heading rules as text_to_pdf since both share
    the ATX heading syntax."""
    return text_to_pdf(md_text)


def _html_escape(s: str) -> str:
    """ReportLab Paragraph parses a tiny HTML subset — we escape the
    characters that could trigger that parser unintentionally."""
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


# ---- DOCX → PDF via LibreOffice -------------------------------------------


def docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX to PDF using LibreOffice in headless mode.

    LibreOffice ships with macOS / most Linux distros and produces
    high-fidelity PDFs (much better than DOCX→HTML→PDF). On systems
    without LibreOffice, this raises RuntimeError and the caller is
    expected to surface it as a 501.
    """
    soffice = _find_libreoffice()
    if not soffice:
        raise RuntimeError(
            "LibreOffice not found on this server — DOCX → PDF conversion "
            "is unavailable. Install with `brew install --cask libreoffice` "
            "or your distro's package manager."
        )

    with tempfile.TemporaryDirectory() as tmp:
        in_path = Path(tmp) / "input.docx"
        in_path.write_bytes(docx_bytes)
        # soffice writes the output PDF next to the input, same basename.
        cmd = [
            soffice, "--headless",
            "--convert-to", "pdf",
            "--outdir", tmp,
            str(in_path),
        ]
        try:
            proc = subprocess.run(
                cmd, check=False, capture_output=True, timeout=120
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError("LibreOffice conversion timed out (>2m)") from exc
        if proc.returncode != 0:
            raise RuntimeError(
                f"LibreOffice failed: {proc.stderr.decode(errors='ignore')[:300]}"
            )
        out_path = Path(tmp) / "input.pdf"
        if not out_path.exists():
            raise RuntimeError("LibreOffice ran but produced no PDF")
        return out_path.read_bytes()


def _find_libreoffice() -> str | None:
    """Look for the soffice / libreoffice binary on common install paths."""
    candidates = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS .dmg install
        "/opt/homebrew/bin/soffice",                              # macOS brew
        "/usr/local/bin/soffice",                                 # macOS brew (Intel)
        "/usr/bin/libreoffice",                                   # Linux apt
        "/usr/bin/soffice",                                       # Linux symlink
        "soffice", "libreoffice",                                 # PATH lookup
    ]
    for c in candidates:
        try:
            r = subprocess.run(
                [c, "--version"], capture_output=True, timeout=5
            )
            if r.returncode == 0:
                return c
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    return None


# ---- Unified entry point ---------------------------------------------------


# ---- Conversion cache ------------------------------------------------------
#
# Conversion can be slow (LibreOffice up to 2min; pdf2docx hundreds of ms per
# page). The same document downloaded by the same user 5 times in the same
# day shouldn't re-run the heavy conversion each time. We cache by
# SHA256(source_bytes) + target_fmt — content-addressed so different source
# documents with identical text don't share cache entries, and unrelated
# documents can never collide.
#
# Cache lives on disk under uploads/conversion_cache/ for persistence across
# server restarts. We add an in-memory LRU on top so common conversions never
# touch the filesystem after the first hit.

import hashlib
from collections import OrderedDict
from threading import Lock

_CONVERSION_CACHE_DIR = (
    Path(__file__).resolve().parents[3] / "uploads" / "conversion_cache"
)
_CONVERSION_CACHE_DIR.mkdir(parents=True, exist_ok=True)
_CACHE_LOCK = Lock()
_CACHE_MAX_ENTRIES = 64
_CACHE_MAX_SOURCE_BYTES = 25 * 1024 * 1024  # don't cache pathologically large inputs
_memory_cache: OrderedDict[str, bytes] = OrderedDict()


def _cache_key(source_bytes: bytes, source_fmt: str, target_fmt: str) -> str:
    h = hashlib.sha256(source_bytes).hexdigest()
    return f"{h}__{source_fmt}__{target_fmt}"


def _cache_get(key: str) -> bytes | None:
    with _CACHE_LOCK:
        if key in _memory_cache:
            # Promote to most-recently-used so the LRU keeps it hot.
            _memory_cache.move_to_end(key)
            return _memory_cache[key]
    disk_path = _CONVERSION_CACHE_DIR / key
    if disk_path.exists():
        try:
            data = disk_path.read_bytes()
        except Exception as exc:
            logger.warning(f"cache read failed {disk_path}: {exc}")
            return None
        with _CACHE_LOCK:
            _memory_cache[key] = data
            _memory_cache.move_to_end(key)
            while len(_memory_cache) > _CACHE_MAX_ENTRIES:
                _memory_cache.popitem(last=False)
        return data
    return None


def _cache_set(key: str, data: bytes) -> None:
    if len(data) > _CACHE_MAX_SOURCE_BYTES:
        # Don't poison the cache with massive outputs.
        return
    with _CACHE_LOCK:
        _memory_cache[key] = data
        _memory_cache.move_to_end(key)
        while len(_memory_cache) > _CACHE_MAX_ENTRIES:
            _memory_cache.popitem(last=False)
    try:
        (_CONVERSION_CACHE_DIR / key).write_bytes(data)
    except Exception as exc:
        logger.warning(f"cache write failed: {exc}")


def convert_document(
    source_bytes: bytes,
    source_fmt: SupportedFormat,
    target_fmt: SupportedFormat,
    *,
    title: str | None = None,
) -> bytes:
    """Convert `source_bytes` from `source_fmt` to `target_fmt`. Bytes-in / bytes-out.

    Routing table (X = direct, ↪ = goes through plain text as intermediate):

                  target →     pdf    docx    txt    md
        source ↓
        pdf                    pass    X      X      X
        docx                   X       pass   X      X
        txt                    X       X      pass   id
        md                     X       X      X      pass

    Raises ValueError for unsupported pairs. Raises RuntimeError for paths
    that require external tools that aren't installed (docx→pdf needs
    LibreOffice).
    """
    src = DocFormat(source_fmt)
    tgt = DocFormat(target_fmt)
    if src == tgt:
        return source_bytes

    # Cache lookup — content-addressed by SHA256 of source + target format.
    # `title` deliberately excluded from the key because most conversions
    # don't use it; the rare PDF/DOCX export with a custom title is fine to
    # re-run (and would otherwise need title in the key, ballooning misses).
    cache_key = _cache_key(source_bytes, src.value, tgt.value)
    cached = _cache_get(cache_key)
    if cached is not None:
        logger.debug(f"conversion cache HIT {src.value}→{tgt.value}")
        return cached

    result: bytes | None = None
    if src == DocFormat.PDF:
        if tgt == DocFormat.DOCX:
            result = pdf_to_docx(source_bytes)
        elif tgt == DocFormat.TXT:
            result = extract_text_from_pdf(source_bytes).encode("utf-8")
        elif tgt == DocFormat.MD:
            result = pdf_to_markdown(source_bytes).encode("utf-8")

    elif src == DocFormat.DOCX:
        if tgt == DocFormat.PDF:
            result = docx_to_pdf(source_bytes)
        elif tgt == DocFormat.TXT:
            result = extract_text_from_docx(source_bytes).encode("utf-8")
        elif tgt == DocFormat.MD:
            result = docx_to_markdown(source_bytes).encode("utf-8")

    elif src == DocFormat.TXT:
        text = source_bytes.decode("utf-8", errors="replace")
        if tgt == DocFormat.PDF:
            result = text_to_pdf(text, title=title)
        elif tgt == DocFormat.DOCX:
            result = text_to_docx(text, title=title)
        elif tgt == DocFormat.MD:
            result = normalize_text(text).encode("utf-8")

    elif src == DocFormat.MD:
        md = source_bytes.decode("utf-8", errors="replace")
        if tgt == DocFormat.PDF:
            result = markdown_to_pdf(md)
        elif tgt == DocFormat.DOCX:
            result = markdown_to_docx(md)
        elif tgt == DocFormat.TXT:
            result = _markdown_strip(md).encode("utf-8")

    if result is None:
        raise ValueError(f"Unsupported conversion: {source_fmt} → {target_fmt}")

    # Cache successful conversion. Skip caching for very large inputs (the
    # _cache_set helper enforces this too on the output side).
    if len(source_bytes) <= _CACHE_MAX_SOURCE_BYTES:
        _cache_set(cache_key, result)
    return result


_MD_HEADING_PREFIX = re.compile(r"^#{1,6}\s+", re.MULTILINE)
_MD_BOLD_ITALIC = re.compile(r"(\*\*|__|\*|_)(.+?)\1")
_MD_INLINE_CODE = re.compile(r"`([^`]+)`")
_MD_LINK = re.compile(r"\[([^\]]+)\]\([^)]+\)")
_MD_LIST_MARKER = re.compile(r"^[-*+]\s+", re.MULTILINE)


def _markdown_strip(md: str) -> str:
    """Remove markdown syntax to produce plain text. Headings and lists keep
    their plain content; emphasis markers and link syntax disappear."""
    s = _MD_HEADING_PREFIX.sub("", md)
    s = _MD_BOLD_ITALIC.sub(r"\2", s)
    s = _MD_INLINE_CODE.sub(r"\1", s)
    s = _MD_LINK.sub(r"\1", s)
    s = _MD_LIST_MARKER.sub("", s)
    return normalize_text(s)


def _mode(values: list[float]) -> float:
    """Return the most common value in `values`, rounded to 1 decimal.

    Used by pdf_to_markdown to estimate the body-text font size. Standard
    statistics.mode raises on ties; we just return the first most common.
    """
    from collections import Counter

    rounded = [round(v, 1) for v in values]
    counter = Counter(rounded)
    return counter.most_common(1)[0][0]
